from __future__ import annotations

from pathlib import Path

from docx import Document

from .base import BaseParser, ParsedContent


class DocxParser(BaseParser):
    extensions = {"docx"}

    def parse(self, path: Path) -> ParsedContent:
        try:
            document = Document(str(path))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            text = "\n".join(paragraphs)
        except Exception:
            paragraphs = []
            result = self.md.convert(str(path))
            text = result.text_content

        truncated = self._truncate(text)
        metadata = {"source": "docx", "paragraphs": len(paragraphs)}
        return ParsedContent(text=truncated, metadata=metadata, page_count=len(document.paragraphs))
